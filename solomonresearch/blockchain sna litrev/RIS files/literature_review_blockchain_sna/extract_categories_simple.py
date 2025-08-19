#!/usr/bin/env python3
"""
Simple Category Content Extractor from Literature Analysis Report
================================================================

This script extracts content for each research category from the final literature 
analysis report using pattern matching, without requiring API calls.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from collections import defaultdict

# Categories to extract
CATEGORIES = [
    "Financial Networks & DeFi",
    "Blockchain-SNA Integration", 
    "Security & Privacy",
    "Blockchain Infrastructure",
    "Network Science & Methods",
    "IoT & Edge Computing",
    "Blockchain Applications",
    "AI & Machine Learning",
    "Supply Chain & Logistics",
    "Economic Models & Game Theory",
    "Literature Reviews",
    "Social Media & Online Networks",
    "Healthcare & Biomedical",
    "Emerging Technologies",
    "Organizational Networks",
    "Governance & Regulation",
    "Related Work",
    "Business Models & Strategy"
]

class SimpleCategoryExtractor:
    def __init__(self, input_file, output_dir):
        self.input_file = input_file
        self.output_dir = output_dir
        self.category_content = defaultdict(list)
        
        # Create output directory
        os.makedirs(self.output_dir, exist_ok=True)
        
    def read_docx_content(self):
        """Read content from the input DOCX file."""
        print(f"📖 Reading document: {self.input_file}")
        doc = Document(self.input_file)
        
        # Extract all paragraphs with their styles
        content_blocks = []
        current_section = ""
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Identify section headers
            if paragraph.style.name.startswith('Heading'):
                current_section = text
                
            content_blocks.append({
                'text': text,
                'style': paragraph.style.name,
                'section': current_section
            })
        
        print(f"✅ Read {len(content_blocks)} content blocks from document")
        return content_blocks
    
    def extract_category_content(self, content_blocks):
        """Extract content for each category based on pattern matching."""
        
        # Sections to focus on
        target_sections = [
            "Key Thematic Patterns",
            "Methodological Trends", 
            "Research Gaps",
            "Significant Research Gaps and Implications",
            "Methodological Evolution",
            "Application Expansion",
            "Integration Sophistication",
            "Societal Impact",
            "Emerging Research Directions"
        ]
        
        print("\n🔍 Extracting content by category...")
        
        for block in content_blocks:
            text = block['text']
            section = block['section']
            
            # Check if we're in a relevant section
            relevant_section = any(target in section for target in target_sections)
            
            # Check each category
            for category in CATEGORIES:
                # Create variations of category name for matching
                category_variations = [
                    category,
                    category.lower(),
                    category.replace('&', 'and'),
                    category.replace(' & ', ' '),
                ]
                
                # Check if category is mentioned in the text
                for variation in category_variations:
                    if variation in text or (relevant_section and variation in text.lower()):
                        # Add the entire paragraph/section that mentions the category
                        self.category_content[category].append({
                            'section': section,
                            'text': text,
                            'style': block['style']
                        })
                        break
        
        # Also extract batch analysis sections that might contain category info
        self.extract_batch_analyses(content_blocks)
    
    def extract_batch_analyses(self, content_blocks):
        """Extract batch analysis sections that contain category-specific insights."""
        
        in_batch_section = False
        current_batch_content = []
        
        for block in content_blocks:
            text = block['text']
            
            # Check if we're entering a batch section
            if "BATCH" in text and "ANALYSIS" in text:
                in_batch_section = True
                current_batch_content = [block]
            elif in_batch_section:
                # Continue collecting batch content
                current_batch_content.append(block)
                
                # Check if batch section is ending
                if "papers)" in text or text.startswith("###"):
                    # Process the batch content
                    self.process_batch_content(current_batch_content)
                    in_batch_section = False
                    current_batch_content = []
    
    def process_batch_content(self, batch_content):
        """Process batch content and assign to relevant categories."""
        
        # Extract categories mentioned in the batch
        batch_text = ' '.join([block['text'] for block in batch_content])
        
        for category in CATEGORIES:
            if category in batch_text or category.lower() in batch_text.lower():
                # Add all batch content to this category
                for block in batch_content:
                    self.category_content[category].append({
                        'section': 'Batch Analysis',
                        'text': block['text'],
                        'style': block['style']
                    })
    
    def create_category_docx(self, category):
        """Create a DOCX file for a specific category."""
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'{category} - Thematic Analysis Extract', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Extracted from: {os.path.basename(self.input_file)}")
        doc.add_paragraph(f"Extraction Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Extracted Sections: {len(self.category_content[category])}")
        doc.add_paragraph()
        
        # Group content by section
        sections = defaultdict(list)
        for content in self.category_content[category]:
            sections[content['section']].append(content)
        
        # Add content organized by section
        for section_name, section_content in sections.items():
            if section_name:
                doc.add_heading(section_name, level=1)
            
            for content in section_content:
                text = content['text']
                style = content['style']
                
                # Apply appropriate formatting
                if 'Heading' in style:
                    level = int(style[-1]) if style[-1].isdigit() else 2
                    doc.add_heading(text, level=min(level, 3))
                elif text.startswith('**') and text.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(text.strip('*')).bold = True
                elif text.startswith('•') or text.startswith('-'):
                    doc.add_paragraph(text, style='List Bullet')
                else:
                    doc.add_paragraph(text)
            
            doc.add_paragraph()  # Add spacing between sections
        
        # Save document
        filename = f"{category.replace('/', '-').replace('&', 'and')}_analysis.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        print(f"💾 Saved: {filename} ({len(self.category_content[category])} sections)")
        return filepath
    
    def process_all_categories(self):
        """Process all categories and create individual DOCX files."""
        print("🚀 Starting category extraction process...")
        print(f"📁 Output directory: {self.output_dir}")
        
        # Read and parse document content
        content_blocks = self.read_docx_content()
        
        # Extract content for all categories
        self.extract_category_content(content_blocks)
        
        # Create DOCX files for each category
        print("\n📝 Creating category-specific documents...")
        files_created = 0
        
        for category in CATEGORIES:
            if self.category_content[category]:
                self.create_category_docx(category)
                files_created += 1
            else:
                print(f"⚠️ No content found for: {category}")
        
        print(f"\n✅ Extraction complete! Created {files_created} files in: {self.output_dir}")
        
        # Create summary file
        self.create_summary_file(files_created)
    
    def create_summary_file(self, files_created):
        """Create a summary file listing extraction results."""
        doc = Document()
        doc.add_heading('Category Extraction Summary', level=0)
        
        doc.add_paragraph(f"Source Document: {self.input_file}")
        doc.add_paragraph(f"Extraction Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Categories Processed: {len(CATEGORIES)}")
        doc.add_paragraph(f"Files Created: {files_created}")
        
        doc.add_heading('Extraction Results', level=1)
        
        # Summary table
        doc.add_paragraph("Category | Sections Found | Status")
        doc.add_paragraph("-" * 50)
        
        for category in CATEGORIES:
            count = len(self.category_content[category])
            status = "✅ Extracted" if count > 0 else "❌ No content"
            filename = f"{category.replace('/', '-').replace('&', 'and')}_analysis.docx"
            
            doc.add_paragraph(f"{category} | {count} sections | {status}")
            if count > 0:
                doc.add_paragraph(f"   → File: {filename}", style='List Bullet')
        
        # Save summary
        summary_path = os.path.join(self.output_dir, "extraction_summary.docx")
        doc.save(summary_path)
        print(f"\n📄 Summary saved: extraction_summary.docx")

def main():
    """Main function to run the extraction process."""
    
    # Configuration
    input_file = "/Users/v/solomonresearch/blockchain sna litrev/RIS files/literature_analysis_report_final.docx"
    output_dir = "/Users/v/solomonresearch/blockchain sna litrev/RIS files/category_extracts"
    
    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"❌ Input file not found: {input_file}")
        return
    
    # Create extractor and process
    extractor = SimpleCategoryExtractor(input_file, output_dir)
    extractor.process_all_categories()

if __name__ == "__main__":
    main()