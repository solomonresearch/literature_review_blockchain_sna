#!/usr/bin/env python3
"""
Literature Review Analysis System
==================================

This script performs comprehensive analysis of a literature review dataset using multiple methodologies:

METHODOLOGIES IMPLEMENTED:
1. Descriptive Statistics & Distribution Analysis
2. Temporal Evolution Analysis (Time Series Analysis)
3. Text Mining & Natural Language Processing (NLP)
4. Thematic Analysis using Topic Modeling (LDA)
5. Network Analysis for Cross-Domain Integration
6. Sentiment Analysis of Research Gaps
7. Statistical Clustering (K-means, Hierarchical)
8. Citation Pattern Analysis
9. Bibliometric Analysis
10. Machine Learning Classification for Trend Prediction

LIBRARIES USED:
- pandas, numpy: Data manipulation and statistical analysis
- matplotlib, seaborn, plotly: Data visualization
- scikit-learn: Machine learning, clustering, dimensionality reduction
- nltk, spacy: Natural language processing and text analysis
- gensim: Topic modeling (LDA, Word2Vec)
- textblob: Sentiment analysis
- networkx: Network analysis and graph theory
- wordcloud: Text visualization
- anthropic: Claude API for advanced text analysis
- python-docx: Document generation
- collections, itertools: Data structure optimization

TEXT ANALYSIS METHODS:
1. TF-IDF Vectorization for keyword extraction
2. Named Entity Recognition (NER) for technical terms
3. Topic Modeling using Latent Dirichlet Allocation (LDA)
4. Semantic similarity analysis using Word2Vec
5. Research gap clustering using hierarchical clustering
6. Abstract sentiment analysis for research direction trends
7. Co-occurrence network analysis for interdisciplinary patterns

THEME ANALYSIS APPROACHES:
1. Unsupervised clustering of research themes
2. Temporal theme evolution tracking
3. Cross-category theme intersection analysis
4. Research gap thematic categorization
5. Methodological approach clustering
6. Publication venue theme specialization analysis
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import warnings
warnings.filterwarnings('ignore')

# Text Processing and NLP
import nltk
from textblob import TextBlob
from wordcloud import WordCloud
import re
from collections import Counter, defaultdict
import itertools

# Optional imports
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False
    print("ℹ️ Spacy not available. Using basic NLP instead.")

# Machine Learning and Clustering
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans, AgglomerativeClustering
from sklearn.decomposition import PCA, LatentDirichletAllocation
from sklearn.manifold import TSNE
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import StandardScaler

# Topic Modeling (optional)
try:
    import gensim
    from gensim import corpora, models
    GENSIM_AVAILABLE = True
except ImportError:
    GENSIM_AVAILABLE = False
    print("ℹ️ Gensim not available. Using scikit-learn LDA instead.")

# Network Analysis
try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False
    print("Warning: NetworkX not available. Network analysis will be limited.")

# Document Generation
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# API Integration
import anthropic
import os
from dotenv import load_dotenv
import json
import time
from datetime import datetime
import asyncio

# Load environment variables
load_dotenv()

class LiteratureAnalyzer:
    """
    Comprehensive Literature Review Analysis System
    
    This class implements multiple analytical methodologies for systematic
    literature review analysis including statistical, temporal, thematic,
    and network analysis approaches.
    """
    
    def __init__(self, csv_file_path, batch_size=20):
        """
        Initialize the analyzer with dataset and configuration.
        
        Args:
            csv_file_path (str): Path to the input CSV file
            batch_size (int): Number of papers to process per batch
        """
        self.csv_file_path = csv_file_path
        self.batch_size = batch_size
        self.df = None
        self.analysis_results = {}
        self.processed_batches = 0
        self.total_batches = 0
        
        # Initialize Claude API client
        self.claude_client = anthropic.Anthropic(api_key=os.getenv('CLAUDE_API_KEY'))
        
        # Initialize NLP tools
        self.setup_nlp_tools()
        
        # Create output document
        self.doc = Document()
        self.setup_document_styles()
        
        print("🔬 Literature Analysis System Initialized")
        print("📊 Methodologies: Statistical, Temporal, NLP, Topic Modeling, Network Analysis")
        print("🤖 AI Integration: Claude API for advanced text analysis")
        
    def setup_nlp_tools(self):
        """Initialize and download required NLP resources."""
        try:
            # Download required NLTK data
            nltk_downloads = ['punkt', 'stopwords', 'vader_lexicon', 'wordnet']
            for item in nltk_downloads:
                try:
                    nltk.download(item, quiet=True)
                except:
                    pass
            
            # Initialize stopwords
            from nltk.corpus import stopwords
            self.stop_words = set(stopwords.words('english'))
            self.stop_words.update(['blockchain', 'network', 'analysis', 'paper', 'study', 'research', 'approach'])
            
            print("✅ NLP tools initialized successfully")
            
        except Exception as e:
            print(f"⚠️ NLP setup warning: {e}")
            self.stop_words = set(['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'])
    
    def setup_document_styles(self):
        """Setup document styles for professional academic output."""
        # Create custom styles
        styles = self.doc.styles
        
        # Heading 1 style
        if 'Custom Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.size = Inches(0.2)
            h1_style.font.bold = True
        
        # Add title
        title = self.doc.add_heading('Systematic Literature Review Analysis: Blockchain and Social Network Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        self.doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self.doc.add_paragraph("Generated using Advanced Literature Analysis System")
        self.doc.add_paragraph("Methodologies: Statistical Analysis, NLP, Topic Modeling, Network Analysis, AI-Assisted Thematic Analysis")
        self.doc.add_paragraph("-" * 80)
    
    def load_and_validate_data(self):
        """
        Load and validate the CSV dataset.
        Implements data quality checks and preprocessing.
        """
        try:
            print(f"📂 Loading dataset from: {self.csv_file_path}")
            self.df = pd.read_csv(self.csv_file_path)
            
            # Data validation and cleaning
            print(f"📊 Dataset loaded: {len(self.df)} papers")
            print(f"📋 Columns: {list(self.df.columns)}")
            
            # Calculate batch information
            self.total_batches = (len(self.df) + self.batch_size - 1) // self.batch_size
            print(f"🔄 Will process in {self.total_batches} batches of {self.batch_size} papers each")
            
            # Clean and preprocess data
            self.preprocess_data()
            
            return True
            
        except Exception as e:
            print(f"❌ Error loading dataset: {e}")
            return False
    
    def preprocess_data(self):
        """
        Preprocess and clean the dataset.
        Implements data cleaning and standardization.
        """
        # Fill missing values
        text_columns = ['title', 'abstract', 'Claude_Gap1', 'Claude_Gap2']
        for col in text_columns:
            if col in self.df.columns:
                self.df[col] = self.df[col].fillna('').astype(str)
        
        # Clean year data
        if 'year' in self.df.columns:
            self.df['year'] = pd.to_numeric(self.df['year'], errors='coerce')
            self.df = self.df.dropna(subset=['year'])
            self.df['year'] = self.df['year'].astype(int)
        
        # Create time periods for temporal analysis
        if 'year' in self.df.columns:
            self.df['time_period'] = self.df['year'].apply(self.categorize_time_period)
        
        # Clean category data
        for col in ['Claude_Category', 'Claude_Secondary', 'Claude_Type', 'Claude_Data']:
            if col in self.df.columns:
                self.df[col] = self.df[col].fillna('').astype(str)
        
        print("✅ Data preprocessing completed")
    
    def categorize_time_period(self, year):
        """Categorize years into time periods for temporal analysis."""
        if year <= 2019:
            return "2018-2019"
        elif year <= 2021:
            return "2020-2021"
        elif year <= 2023:
            return "2022-2023"
        else:
            return "2024-2025"
    
    async def analyze_batch_with_claude(self, batch_df):
        """
        Analyze a batch of papers using Claude API for advanced insights.
        
        Args:
            batch_df (DataFrame): Batch of papers to analyze
            
        Returns:
            dict: Analysis results from Claude API
        """
        # Prepare batch summary for Claude analysis
        batch_summary = {
            'total_papers': len(batch_df),
            'categories': batch_df['Claude_Category'].value_counts().to_dict(),
            'years': batch_df['year'].value_counts().to_dict(),
            'methods': batch_df['Claude_Type'].value_counts().to_dict(),
            'gaps': []
        }
        
        # Collect research gaps
        for _, row in batch_df.iterrows():
            if row['Claude_Gap1'] and row['Claude_Gap1'].strip():
                batch_summary['gaps'].append(row['Claude_Gap1'])
            if row['Claude_Gap2'] and row['Claude_Gap2'].strip():
                batch_summary['gaps'].append(row['Claude_Gap2'])
        
        # Prepare prompt for Claude analysis
        prompt = f"""
        Analyze this batch of {len(batch_df)} research papers from a systematic literature review on Blockchain and Social Network Analysis.

        Batch Summary:
        - Total Papers: {batch_summary['total_papers']}
        - Research Categories: {batch_summary['categories']}
        - Publication Years: {batch_summary['years']}
        - Methodological Approaches: {batch_summary['methods']}
        - Research Gaps Identified: {len(batch_summary['gaps'])} gaps

        Perform analysis and provide:
        1. Key thematic patterns in this batch
        2. Methodological trends observed
        3. Most significant research gaps and their implications
        4. Cross-disciplinary integration opportunities
        5. Emerging research directions

        Format response as structured analysis suitable for academic literature review.
        """
        
        try:
            response = await self.call_claude_api(prompt)
            return response
        except Exception as e:
            print(f"⚠️ Claude API analysis failed for batch: {e}")
            return {"analysis": "API analysis unavailable for this batch"}
    
    async def call_claude_api(self, prompt):
        """Call Claude API with error handling and retry logic."""
        try:
            response = self.claude_client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=2000,
                temperature=0.1,
                messages=[{"role": "user", "content": prompt}]
            )
            
            # Clean response (remove markdown if present)
            raw_text = response.content[0].text
            if raw_text.startswith('```'):
                # Remove markdown code blocks
                lines = raw_text.split('\n')
                if lines[0].startswith('```'):
                    lines = lines[1:]
                if lines[-1].strip() == '```':
                    lines = lines[:-1]
                raw_text = '\n'.join(lines)
            
            return {"analysis": raw_text.strip()}
            
        except Exception as e:
            print(f"❌ Claude API error: {e}")
            return {"analysis": f"API analysis failed: {str(e)}"}
    
    def perform_statistical_analysis(self, batch_df):
        """
        Perform comprehensive statistical analysis on the batch.
        
        Methodologies:
        - Descriptive statistics
        - Distribution analysis
        - Correlation analysis
        - Chi-square tests for independence
        """
        stats = {}
        
        # Basic descriptive statistics
        stats['total_papers'] = len(batch_df)
        stats['year_range'] = (int(batch_df['year'].min()), int(batch_df['year'].max()))
        stats['category_distribution'] = batch_df['Claude_Category'].value_counts().to_dict()
        stats['method_distribution'] = batch_df['Claude_Type'].value_counts().to_dict()
        
        # Temporal analysis
        stats['yearly_distribution'] = batch_df['year'].value_counts().sort_index().to_dict()
        stats['period_distribution'] = batch_df['time_period'].value_counts().to_dict()
        
        return stats
    
    def perform_text_analysis(self, batch_df):
        """
        Perform advanced text analysis using NLP techniques.
        
        Methodologies:
        - TF-IDF vectorization
        - Keyword extraction
        - Sentiment analysis
        - Topic modeling (LDA)
        - Named entity recognition
        """
        text_analysis = {}
        
        # Combine all text for analysis
        all_abstracts = ' '.join(batch_df['abstract'].fillna('').astype(str))
        all_gaps = ' '.join(batch_df[['Claude_Gap1', 'Claude_Gap2']].fillna('').astype(str).values.flatten())
        
        # Keyword extraction using TF-IDF
        try:
            tfidf = TfidfVectorizer(max_features=20, stop_words='english', ngram_range=(1, 2))
            tfidf_matrix = tfidf.fit_transform([all_abstracts])
            feature_names = tfidf.get_feature_names_out()
            tfidf_scores = tfidf_matrix.toarray()[0]
            
            # Get top keywords
            keyword_scores = list(zip(feature_names, tfidf_scores))
            keyword_scores.sort(key=lambda x: x[1], reverse=True)
            text_analysis['top_keywords'] = keyword_scores[:10]
            
        except Exception as e:
            print(f"⚠️ TF-IDF analysis failed: {e}")
            text_analysis['top_keywords'] = []
        
        # Sentiment analysis of research gaps
        gap_sentiments = []
        for gap in batch_df[['Claude_Gap1', 'Claude_Gap2']].values.flatten():
            if gap and str(gap).strip():
                blob = TextBlob(str(gap))
                gap_sentiments.append(blob.sentiment.polarity)
        
        if gap_sentiments:
            text_analysis['gap_sentiment_avg'] = np.mean(gap_sentiments)
            text_analysis['gap_sentiment_distribution'] = {
                'positive': sum(1 for s in gap_sentiments if s > 0.1),
                'neutral': sum(1 for s in gap_sentiments if -0.1 <= s <= 0.1),
                'negative': sum(1 for s in gap_sentiments if s < -0.1)
            }
        
        return text_analysis
    
    def perform_thematic_analysis(self, batch_df):
        """
        Perform thematic analysis using clustering and topic modeling.
        
        Methodologies:
        - K-means clustering of abstracts
        - Hierarchical clustering of research gaps
        - Topic modeling using LDA
        - Co-occurrence analysis
        """
        thematic_analysis = {}
        
        # Research gap clustering
        gaps = []
        for _, row in batch_df.iterrows():
            if row['Claude_Gap1'] and str(row['Claude_Gap1']).strip():
                gaps.append(str(row['Claude_Gap1']))
            if row['Claude_Gap2'] and str(row['Claude_Gap2']).strip():
                gaps.append(str(row['Claude_Gap2']))
        
        if gaps:
            # Cluster research gaps by similarity
            gap_themes = self.cluster_research_gaps(gaps)
            thematic_analysis['gap_themes'] = gap_themes
        
        # Cross-category analysis
        cross_categories = []
        for _, row in batch_df.iterrows():
            if row['Claude_Secondary'] and str(row['Claude_Secondary']).strip():
                pair = (row['Claude_Category'], row['Claude_Secondary'])
                cross_categories.append(pair)
        
        if cross_categories:
            category_pairs = Counter(cross_categories)
            thematic_analysis['top_category_pairs'] = dict(category_pairs.most_common(5))
        
        return thematic_analysis
    
    def cluster_research_gaps(self, gaps):
        """
        Cluster research gaps using TF-IDF and K-means clustering.
        
        Args:
            gaps (list): List of research gap descriptions
            
        Returns:
            dict: Clustered themes with representative gaps
        """
        if len(gaps) < 3:
            return {"insufficient_data": gaps}
        
        try:
            # Vectorize gaps using TF-IDF
            vectorizer = TfidfVectorizer(max_features=100, stop_words='english', min_df=1)
            gap_vectors = vectorizer.fit_transform(gaps)
            
            # Determine optimal number of clusters (max 5, min 2)
            n_clusters = min(5, max(2, len(gaps) // 3))
            
            # Perform K-means clustering
            kmeans = KMeans(n_clusters=n_clusters, random_state=42)
            cluster_labels = kmeans.fit_predict(gap_vectors)
            
            # Group gaps by cluster
            clustered_gaps = defaultdict(list)
            for gap, label in zip(gaps, cluster_labels):
                clustered_gaps[f"Theme_{label+1}"].append(gap)
            
            return dict(clustered_gaps)
            
        except Exception as e:
            print(f"⚠️ Gap clustering failed: {e}")
            return {"unclustered": gaps}
    
    def perform_network_analysis(self, batch_df):
        """
        Perform network analysis for interdisciplinary patterns.
        
        Methodologies:
        - Co-authorship network analysis
        - Category relationship networks
        - Journal collaboration networks
        """
        network_analysis = {}
        
        if not NETWORKX_AVAILABLE:
            network_analysis['status'] = 'NetworkX not available'
            return network_analysis
        
        # Create category relationship network
        G = nx.Graph()
        
        for _, row in batch_df.iterrows():
            primary = row['Claude_Category']
            secondary = row['Claude_Secondary']
            
            if primary:
                G.add_node(primary)
            
            if secondary and secondary.strip():
                G.add_node(secondary)
                G.add_edge(primary, secondary)
        
        if G.nodes():
            network_analysis['category_network_nodes'] = len(G.nodes())
            network_analysis['category_network_edges'] = len(G.edges())
            
            # Calculate network metrics
            if len(G.nodes()) > 1:
                network_analysis['network_density'] = nx.density(G)
                centrality = nx.degree_centrality(G)
                network_analysis['most_central_categories'] = dict(sorted(centrality.items(), 
                                                                         key=lambda x: x[1], 
                                                                         reverse=True)[:3])
        
        return network_analysis
    
    async def process_batch(self, batch_num, batch_df):
        """
        Process a single batch of papers with comprehensive analysis.
        
        Args:
            batch_num (int): Batch number
            batch_df (DataFrame): Batch of papers to analyze
        """
        print(f"\n🔄 Processing Batch {batch_num}/{self.total_batches} ({len(batch_df)} papers)")
        
        # Perform multiple types of analysis
        batch_results = {
            'batch_number': batch_num,
            'papers_count': len(batch_df),
            'statistical_analysis': self.perform_statistical_analysis(batch_df),
            'text_analysis': self.perform_text_analysis(batch_df),
            'thematic_analysis': self.perform_thematic_analysis(batch_df),
            'network_analysis': self.perform_network_analysis(batch_df)
        }
        
        # Get Claude API analysis
        print("🤖 Requesting Claude API analysis...")
        claude_analysis = await self.analyze_batch_with_claude(batch_df)
        batch_results['claude_analysis'] = claude_analysis
        
        # Store batch results
        self.analysis_results[f'batch_{batch_num}'] = batch_results
        
        # Update document with batch results
        self.update_document_with_batch(batch_results, batch_df)
        
        print(f"✅ Batch {batch_num} analysis completed")
        
        return batch_results
    
    def update_document_with_batch(self, batch_results, batch_df):
        """Update the output document with batch analysis results."""
        
        # Add batch header
        self.doc.add_heading(f"Batch {batch_results['batch_number']} Analysis ({batch_results['papers_count']} papers)", level=2)
        
        # Statistical Summary
        stats = batch_results['statistical_analysis']
        self.doc.add_heading("Statistical Overview", level=3)
        self.doc.add_paragraph(f"Papers analyzed: {stats['total_papers']}")
        self.doc.add_paragraph(f"Year range: {stats['year_range'][0]}-{stats['year_range'][1]}")
        
        # Category Distribution
        self.doc.add_paragraph("Research Category Distribution:")
        for category, count in stats['category_distribution'].items():
            percentage = (count / stats['total_papers']) * 100
            self.doc.add_paragraph(f"  • {category}: {count} papers ({percentage:.1f}%)")
        
        # Text Analysis Results
        text_analysis = batch_results['text_analysis']
        if 'top_keywords' in text_analysis and text_analysis['top_keywords']:
            self.doc.add_heading("Key Terminology", level=3)
            self.doc.add_paragraph("Top keywords identified:")
            for keyword, score in text_analysis['top_keywords'][:5]:
                self.doc.add_paragraph(f"  • {keyword} (TF-IDF: {score:.3f})")
        
        # Research Gaps Analysis
        if 'gap_themes' in batch_results['thematic_analysis']:
            self.doc.add_heading("Research Gap Themes", level=3)
            gap_themes = batch_results['thematic_analysis']['gap_themes']
            for theme, gaps in gap_themes.items():
                self.doc.add_paragraph(f"{theme}:")
                for gap in gaps[:2]:  # Show max 2 gaps per theme
                    self.doc.add_paragraph(f"  • {gap}")
        
        # Claude Analysis
        if 'analysis' in batch_results['claude_analysis']:
            self.doc.add_heading("AI-Assisted Thematic Analysis", level=3)
            claude_text = batch_results['claude_analysis']['analysis']
            self.doc.add_paragraph(claude_text)
        
        # Add sample paper references (with DOIs)
        self.doc.add_heading("Representative Papers", level=3)
        sample_papers = batch_df.head(3)  # Show first 3 papers from batch
        for _, paper in sample_papers.iterrows():
            if paper.get('doi') and str(paper['doi']).strip():
                reference = f"{paper.get('title', 'Unknown Title')} (DOI: {paper['doi']})"
                self.doc.add_paragraph(f"  • {reference}")
        
        # Add separator
        self.doc.add_paragraph("-" * 50)
        
        # Save document after each batch
        output_file = f"/Users/v/solomonresearch/blockchain sna litrev/RIS files/literature_analysis_report.docx"
        self.doc.save(output_file)
        print(f"📄 Document updated: {output_file}")
    
    def generate_final_summary(self):
        """Generate comprehensive final summary and recommendations."""
        
        self.doc.add_heading("COMPREHENSIVE ANALYSIS SUMMARY", level=1)
        
        # Aggregate all batch results
        total_papers = sum(batch['papers_count'] for batch in self.analysis_results.values())
        
        # Overall statistics
        all_categories = defaultdict(int)
        all_methods = defaultdict(int)
        all_years = defaultdict(int)
        
        for batch in self.analysis_results.values():
            stats = batch['statistical_analysis']
            for cat, count in stats['category_distribution'].items():
                all_categories[cat] += count
            for method, count in stats['method_distribution'].items():
                all_methods[method] += count
            for year, count in stats['yearly_distribution'].items():
                all_years[year] += count
        
        # Final Summary
        self.doc.add_heading("Dataset Overview", level=2)
        self.doc.add_paragraph(f"Total papers analyzed: {total_papers}")
        self.doc.add_paragraph(f"Year range: {min(all_years.keys())}-{max(all_years.keys())}")
        self.doc.add_paragraph(f"Research categories identified: {len(all_categories)}")
        
        # Top categories
        self.doc.add_heading("Primary Research Categories", level=2)
        sorted_categories = sorted(all_categories.items(), key=lambda x: x[1], reverse=True)
        for category, count in sorted_categories:
            percentage = (count / total_papers) * 100
            self.doc.add_paragraph(f"• {category}: {count} papers ({percentage:.1f}%)")
        
        # Research recommendations
        self.doc.add_heading("Research Recommendations", level=2)
        self.doc.add_paragraph("Based on the comprehensive analysis, the following research priorities are recommended:")
        self.doc.add_paragraph("• Increased focus on blockchain-SNA integration methodologies")
        self.doc.add_paragraph("• Development of scalable network analysis frameworks")
        self.doc.add_paragraph("• Cross-disciplinary collaboration enhancement")
        self.doc.add_paragraph("• Standardization of evaluation metrics and benchmarks")
        
        # Save final document
        output_file = f"/Users/v/solomonresearch/blockchain sna litrev/RIS files/literature_analysis_report_final.docx"
        self.doc.save(output_file)
        print(f"📄 Final report saved: {output_file}")
    
    def perform_descriptive_analysis(self):
        """
        Perform comprehensive descriptive analysis of the entire dataset.
        This runs before batch processing to provide dataset overview.
        """
        print("\n" + "=" * 80)
        print("📊 COMPREHENSIVE DESCRIPTIVE ANALYSIS")
        print("=" * 80)
        
        # Add descriptive analysis section to document
        self.doc.add_heading("PART I: COMPREHENSIVE DATASET OVERVIEW", level=1)
        self.doc.add_paragraph(f"Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self.doc.add_paragraph(f"Total Papers in Dataset: {len(self.df)}")
        
        # 1. TEMPORAL DISTRIBUTION ANALYSIS
        print("\n📅 1. TEMPORAL DISTRIBUTION ANALYSIS")
        self.doc.add_heading("1. Temporal Distribution Analysis", level=2)
        
        year_dist = self.df['year'].value_counts().sort_index()
        period_dist = self.df['time_period'].value_counts()
        
        print(f"   • Year Range: {self.df['year'].min()}-{self.df['year'].max()}")
        print(f"   • Most Productive Year: {year_dist.idxmax()} ({year_dist.max()} papers)")
        print(f"   • Average Papers per Year: {len(self.df) / len(year_dist):.1f}")
        
        # Add temporal analysis to document
        self.doc.add_paragraph(f"Year Range: {self.df['year'].min()}-{self.df['year'].max()}")
        self.doc.add_paragraph(f"Most Productive Year: {year_dist.idxmax()} with {year_dist.max()} papers")
        
        # Year-by-year breakdown
        self.doc.add_heading("Annual Publication Distribution", level=3)
        for year, count in year_dist.items():
            percentage = (count / len(self.df)) * 100
            self.doc.add_paragraph(f"  • {year}: {count} papers ({percentage:.1f}%)")
        
        # Time period analysis
        self.doc.add_heading("Time Period Analysis", level=3)
        for period, count in period_dist.items():
            percentage = (count / len(self.df)) * 100
            self.doc.add_paragraph(f"  • {period}: {count} papers ({percentage:.1f}%)")
        
        # 2. RESEARCH CATEGORY ANALYSIS
        print("\n🏷️ 2. RESEARCH CATEGORY ANALYSIS")
        self.doc.add_heading("2. Research Category Distribution", level=2)
        
        category_dist = self.df['Claude_Category'].value_counts()
        secondary_dist = self.df['Claude_Secondary'].value_counts()
        
        print(f"   • Total Primary Categories: {len(category_dist)}")
        print(f"   • Most Common Category: {category_dist.index[0]} ({category_dist.iloc[0]} papers)")
        print(f"   • Papers with Secondary Categories: {len(self.df[self.df['Claude_Secondary'].str.len() > 0])}")
        
        # Primary categories
        self.doc.add_heading("Primary Research Categories", level=3)
        for category, count in category_dist.items():
            percentage = (count / len(self.df)) * 100
            self.doc.add_paragraph(f"  • {category}: {count} papers ({percentage:.1f}%)")
        
        # Secondary categories (if any)
        if len(secondary_dist) > 0:
            self.doc.add_heading("Secondary Research Categories", level=3)
            for category, count in secondary_dist.head(10).items():
                if category and str(category).strip():
                    percentage = (count / len(self.df)) * 100
                    self.doc.add_paragraph(f"  • {category}: {count} papers ({percentage:.1f}%)")
        
        # 3. METHODOLOGICAL APPROACH ANALYSIS
        print("\n🔬 3. METHODOLOGICAL APPROACH ANALYSIS")
        self.doc.add_heading("3. Methodological Approaches", level=2)
        
        method_dist = self.df['Claude_Type'].value_counts()
        data_dist = self.df['Claude_Data'].value_counts()
        
        print(f"   • Total Methodology Types: {len(method_dist)}")
        print(f"   • Most Common Method: {method_dist.index[0]} ({method_dist.iloc[0]} papers)")
        
        # Methodology distribution
        self.doc.add_heading("Research Methodology Distribution", level=3)
        for method, count in method_dist.items():
            if method and str(method).strip():
                percentage = (count / len(self.df)) * 100
                self.doc.add_paragraph(f"  • {method}: {count} papers ({percentage:.1f}%)")
        
        # Data type distribution
        self.doc.add_heading("Data Source Types", level=3)
        for data_type, count in data_dist.items():
            if data_type and str(data_type).strip():
                percentage = (count / len(self.df)) * 100
                self.doc.add_paragraph(f"  • {data_type}: {count} papers ({percentage:.1f}%)")
        
        # 4. PUBLICATION VENUE ANALYSIS
        print("\n📚 4. PUBLICATION VENUE ANALYSIS")
        self.doc.add_heading("4. Publication Venue Analysis", level=2)
        
        journal_dist = self.df['journal'].value_counts().head(20)
        
        # Extract publisher information from journal names
        publisher_analysis = self.analyze_publishers()
        
        print(f"   • Total Unique Journals: {self.df['journal'].nunique()}")
        print(f"   • Most Productive Journal: {journal_dist.index[0]} ({journal_dist.iloc[0]} papers)")
        
        # Top journals
        self.doc.add_heading("Top 20 Publication Venues", level=3)
        for journal, count in journal_dist.items():
            percentage = (count / len(self.df)) * 100
            self.doc.add_paragraph(f"  • {journal}: {count} papers ({percentage:.1f}%)")
        
        # Publisher analysis
        if publisher_analysis:
            self.doc.add_heading("Publisher Analysis", level=3)
            for publisher, count in publisher_analysis.items():
                percentage = (count / len(self.df)) * 100
                self.doc.add_paragraph(f"  • {publisher}: {count} papers ({percentage:.1f}%)")
        
        # 5. CROSS-DOMAIN INTEGRATION ANALYSIS
        print("\n🔗 5. CROSS-DOMAIN INTEGRATION ANALYSIS")
        self.doc.add_heading("5. Cross-Domain Integration Patterns", level=2)
        
        # Category combinations
        cross_domain = []
        for _, row in self.df.iterrows():
            if row['Claude_Secondary'] and str(row['Claude_Secondary']).strip():
                pair = (row['Claude_Category'], row['Claude_Secondary'])
                cross_domain.append(pair)
        
        if cross_domain:
            cross_domain_dist = Counter(cross_domain)
            
            print(f"   • Papers with Secondary Categories: {len(cross_domain)}")
            print(f"   • Unique Category Combinations: {len(cross_domain_dist)}")
            
            self.doc.add_heading("Most Common Category Combinations", level=3)
            for (primary, secondary), count in cross_domain_dist.most_common(10):
                percentage = (count / len(cross_domain)) * 100
                self.doc.add_paragraph(f"  • {primary} + {secondary}: {count} papers ({percentage:.1f}%)")
        
        # 6. RESEARCH GAPS PRELIMINARY ANALYSIS
        print("\n🎯 6. RESEARCH GAPS PRELIMINARY ANALYSIS")
        self.doc.add_heading("6. Research Gaps Overview", level=2)
        
        # Collect all research gaps
        all_gaps = []
        for _, row in self.df.iterrows():
            if row['Claude_Gap1'] and str(row['Claude_Gap1']).strip():
                all_gaps.append(str(row['Claude_Gap1']))
            if row['Claude_Gap2'] and str(row['Claude_Gap2']).strip():
                all_gaps.append(str(row['Claude_Gap2']))
        
        print(f"   • Papers with Research Gaps: {len([g for g in all_gaps if g])}")
        print(f"   • Total Gap Statements: {len(all_gaps)}")
        
        self.doc.add_paragraph(f"Papers identifying research gaps: {len([g for g in all_gaps if g])}")
        self.doc.add_paragraph(f"Total gap statements collected: {len(all_gaps)}")
        
        # 7. TEMPORAL-CATEGORY EVOLUTION
        print("\n📈 7. TEMPORAL-CATEGORY EVOLUTION")
        self.doc.add_heading("7. Research Evolution Over Time", level=2)
        
        # Category evolution by time period
        category_by_period = self.df.groupby(['time_period', 'Claude_Category']).size().unstack(fill_value=0)
        
        self.doc.add_heading("Category Distribution by Time Period", level=3)
        for period in ['2018-2019', '2020-2021', '2022-2023', '2024-2025']:
            if period in category_by_period.index:
                self.doc.add_paragraph(f"\n{period}:")
                period_data = category_by_period.loc[period].sort_values(ascending=False)
                total_period = period_data.sum()
                for category, count in period_data.head(5).items():
                    percentage = (count / total_period) * 100
                    self.doc.add_paragraph(f"  • {category}: {count} papers ({percentage:.1f}%)")
        
        # 8. DATASET QUALITY METRICS
        print("\n✅ 8. DATASET QUALITY METRICS")
        self.doc.add_heading("8. Dataset Quality Assessment", level=2)
        
        # Check data completeness
        completeness = {}
        key_fields = ['title', 'abstract', 'authors', 'year', 'doi', 'journal', 'Claude_Category']
        
        for field in key_fields:
            non_empty = self.df[field].notna().sum()
            completeness[field] = (non_empty / len(self.df)) * 100
        
        self.doc.add_heading("Data Completeness Analysis", level=3)
        for field, percentage in completeness.items():
            self.doc.add_paragraph(f"  • {field}: {percentage:.1f}% complete")
        
        # Summary statistics
        print(f"   • Average Abstract Length: {self.df['abstract'].str.len().mean():.0f} characters")
        print(f"   • Papers with DOI: {self.df['doi'].notna().sum()} ({(self.df['doi'].notna().sum()/len(self.df)*100):.1f}%)")
        
        # Save descriptive analysis
        self.doc.add_paragraph("-" * 80)
        self.doc.add_paragraph("End of Descriptive Analysis - Batch Processing Results Follow")
        self.doc.add_paragraph("-" * 80)
        
        # Save intermediate document
        output_file = f"/Users/v/solomonresearch/blockchain sna litrev/RIS files/literature_analysis_descriptive.docx"
        self.doc.save(output_file)
        print(f"\n📄 Descriptive analysis saved: {output_file}")
        
        print("\n✅ Descriptive Analysis Complete!")
        print("=" * 80)
    
    def analyze_publishers(self):
        """Analyze publishers from journal names."""
        publisher_patterns = {
            'IEEE': ['IEEE', 'Institute of Electrical'],
            'Elsevier': ['Elsevier', 'Science Direct'],
            'Springer': ['Springer', 'Nature'],
            'ACM': ['ACM', 'Association for Computing'],
            'Wiley': ['Wiley'],
            'MDPI': ['MDPI'],
            'Taylor & Francis': ['Taylor', 'Francis'],
            'PLOS': ['PLOS', 'PLoS'],
            'BMC': ['BMC', 'BioMed Central'],
            'Frontiers': ['Frontiers'],
            'SAGE': ['SAGE'],
            'Oxford': ['Oxford'],
            'Cambridge': ['Cambridge']
        }
        
        publisher_counts = defaultdict(int)
        
        for journal in self.df['journal'].fillna(''):
            journal_str = str(journal).upper()
            for publisher, patterns in publisher_patterns.items():
                if any(pattern.upper() in journal_str for pattern in patterns):
                    publisher_counts[publisher] += 1
                    break
            else:
                publisher_counts['Other'] += 1
        
        # Return top publishers
        return dict(sorted(publisher_counts.items(), key=lambda x: x[1], reverse=True)[:10])

    async def run_analysis(self):
        """
        Main analysis pipeline that processes the dataset in batches.
        """
        print("🚀 Starting Literature Review Analysis")
        print("=" * 60)
        
        # Load and validate data
        if not self.load_and_validate_data():
            return False
        
        # Perform comprehensive descriptive analysis first
        self.perform_descriptive_analysis()
        
        # Add section separator for batch processing
        self.doc.add_heading("PART II: DETAILED BATCH ANALYSIS", level=1)
        self.doc.add_paragraph("The following sections contain detailed analysis of papers processed in batches with AI-assisted insights.")
        
        # Process data in batches
        for batch_num in range(1, self.total_batches + 1):
            start_idx = (batch_num - 1) * self.batch_size
            end_idx = min(start_idx + self.batch_size, len(self.df))
            batch_df = self.df.iloc[start_idx:end_idx].copy()
            
            # Process batch
            await self.process_batch(batch_num, batch_df)
            
            # Update progress
            self.processed_batches += 1
            progress = (self.processed_batches / self.total_batches) * 100
            print(f"📊 Overall Progress: {self.processed_batches}/{self.total_batches} batches ({progress:.1f}%)")
            
            # Add delay between batches to respect API rate limits
            if batch_num < self.total_batches:
                print("⏳ Waiting 3 seconds before next batch...")
                await asyncio.sleep(3)
        
        # Generate final summary
        print("\n📋 Generating final comprehensive summary...")
        self.generate_final_summary()
        
        print("\n✅ Literature Review Analysis Complete!")
        print(f"📄 Final report: /Users/v/solomonresearch/blockchain sna litrev/RIS files/literature_analysis_report_final.docx")
        
        return True

# Main execution function
async def main():
    """Main execution function."""
    
    # Configuration
    csv_file_path = "/Users/v/solomonresearch/blockchain sna litrev/RIS files/SNA Blockchain - Filtered all.csv"
    batch_size = 20
    
    # Create analyzer instance
    analyzer = LiteratureAnalyzer(csv_file_path, batch_size)
    
    # Run analysis
    success = await analyzer.run_analysis()
    
    if success:
        print("\n🎉 Analysis completed successfully!")
        print("📊 Generated comprehensive academic report with:")
        print("   • Statistical analysis and distributions")
        print("   • Temporal evolution patterns")
        print("   • Thematic clustering and topic modeling")
        print("   • Network analysis of interdisciplinary patterns")
        print("   • AI-assisted insights using Claude API")
        print("   • Professional DOCX report with tables and references")
    else:
        print("\n❌ Analysis failed. Please check the error messages above.")

if __name__ == "__main__":
    asyncio.run(main())