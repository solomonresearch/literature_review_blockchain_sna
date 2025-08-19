#!/usr/bin/env python3
"""
Quick verification script to check extracted category content
"""

from docx import Document
import os

def verify_extracted_content():
    """Verify the quality of extracted content."""
    
    extract_dir = "/Users/v/solomonresearch/blockchain sna litrev/RIS files/category_extracts"
    
    # Check a few sample files
    sample_files = [
        "Financial Networks and DeFi_analysis.docx",
        "Blockchain-SNA Integration_analysis.docx", 
        "AI and Machine Learning_analysis.docx"
    ]
    
    for filename in sample_files:
        filepath = os.path.join(extract_dir, filename)
        if os.path.exists(filepath):
            print(f"\n{'='*60}")
            print(f"📄 CONTENT PREVIEW: {filename}")
            print('='*60)
            
            doc = Document(filepath)
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and paragraph_count < 10:  # Show first 10 paragraphs
                    print(f"{paragraph_count + 1:2d}. {text[:100]}...")
                    paragraph_count += 1
                elif paragraph_count >= 10:
                    break
            
            print(f"\nTotal paragraphs in document: {len([p for p in doc.paragraphs if p.text.strip()])}")
        else:
            print(f"❌ File not found: {filename}")

if __name__ == "__main__":
    verify_extracted_content()